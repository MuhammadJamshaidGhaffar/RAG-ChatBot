"""
File handling service layer
"""

import os
import uuid
import sys
import csv
import time
import json
from pathlib import Path
from datetime import datetime
from fastapi import UploadFile
from fastapi.responses import JSONResponse
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from docx.oxml.shared import qn
import os
# Add the project root to the Python path
project_root = Path(__file__).parent.parent.parent
sys.path.append(str(project_root))

# Import database utilities
sys.path.append(str(project_root / "database"))
from mongo_client import get_mongo_client
from utils.constants import IMAGES_COLLECTION, VIDEOS_COLLECTION, EXTRAS_COLLECTION, CHUNK_SIZE, CHUNK_OVERLAP

from utils.vector_db import get_pinecone_vector_store, get_pinecone_stats, list_uploaded_files, delete_file_from_pinecone, add_file_to_database


class ProgressTrackingVectorStore:
    """Wrapper around vector store to provide progress updates during embedding generation."""
    
    def __init__(self, vector_store, upload_id, progress_callback):
        self.vector_store = vector_store
        self.upload_id = upload_id
        self.progress_callback = progress_callback
    
    def add_documents(self, documents, ids):
        """Add documents with progress tracking for each individual document."""
        total_docs = len(documents)
        
        for i, (doc, doc_id) in enumerate(zip(documents, ids)):
            # Update progress for individual chunk
            if self.progress_callback:
                self.progress_callback(self.upload_id, {
                    "message": f"Generating embedding for chunk {i+1}/{total_docs} in current batch...",
                    "updated_at": datetime.now().isoformat()
                })
            
            # Add individual document
            self.vector_store.add_documents(documents=[doc], ids=[doc_id])
            
            # Small delay to allow progress updates to be visible
            time.sleep(0.1)


def add_documents_to_vector_store_with_batching(vector_store, documents, metadatas=None, batch_size=2, max_retries=5):
    """
    Add documents to the Pinecone vector store in batches with exponential backoff.
    
    Args:
        vector_store: PineconeVectorStore instance
        documents: List of document strings to add
        metadatas: Optional list of metadata dictionaries
        batch_size: Number of chunks to process at once (default: 2)
        max_retries: Maximum number of retries for each batch (default: 5)
        
    Returns:
        dict: Results including success count, failed batches, and retry info
    """
    if not isinstance(documents, list):
        raise ValueError("Documents should be a list of strings.")
    
    if metadatas is None:
        metadatas = [{} for _ in documents]
    
    if len(documents) != len(metadatas):
        raise ValueError("Length of documents and metadatas must match.")
    
    print(f"DEBUG: Starting batch processing of {len(documents)} documents")
    
    # Split documents into chunks first
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    
    print(f"DEBUG: Splitting {len(documents)} documents into chunks")
    all_chunks = text_splitter.create_documents(documents, metadatas=metadatas)
    total_chunks = len(all_chunks)
    
    print(f"DEBUG: Created {total_chunks} chunks, processing in batches of {batch_size}")
    
    # Process chunks in batches
    successful_chunks = 0
    failed_batches = []
    current_batch_start = 0
    
    try:
        while current_batch_start < total_chunks:
            batch_end = min(current_batch_start + batch_size, total_chunks)
            batch_chunks = all_chunks[current_batch_start:batch_end]
            batch_number = (current_batch_start // batch_size) + 1
            
            print(f"DEBUG: Processing batch {batch_number} - chunks {current_batch_start + 1} to {batch_end}")
            
            # Try to process this batch with exponential backoff
            retry_count = 0
            batch_success = False
            
            while retry_count < max_retries and not batch_success:
                try:
                    # Generate UUIDs for this batch
                    batch_uuids = [str(uuid.uuid4()) for _ in range(len(batch_chunks))]
                    
                    print(batch_chunks)
                    # Add documents to vector store (this is where embeddings are generated)
                    vector_store.add_documents(documents=batch_chunks, ids=batch_uuids)
                    
                    successful_chunks += len(batch_chunks)
                    batch_success = True
                    print(f"DEBUG: Successfully processed batch {batch_number} ({len(batch_chunks)} chunks)")
                    
                except Exception as e:
                    retry_count += 1
                    error_msg = str(e)
                    
                    # Check if it's a quota/rate limit error
                    is_quota_error = any(keyword in error_msg.lower() for keyword in [
                        'quota', 'rate limit', 'too many requests', 'resource exhausted',
                        'exceeded', 'limit', 'throttle', 'usage'
                    ])
                    
                    if is_quota_error and retry_count < max_retries:
                        # Exponential backoff: 2^retry_count seconds (2, 4, 8, 16, 32)
                        wait_time = 2 ** retry_count
                        print(f"DEBUG: Quota/rate limit error in batch {batch_number}, retry {retry_count}/{max_retries}. "
                              f"Waiting {wait_time} seconds before retry...")
                        print(f"DEBUG: Error details: {error_msg}")
                        
                        time.sleep(wait_time)
                    else:
                        print(f"DEBUG: Failed to process batch {batch_number} after {retry_count} retries: {error_msg}")
                        failed_batches.append({
                            "batch_number": batch_number,
                            "start_chunk": current_batch_start,
                            "end_chunk": batch_end,
                            "error": error_msg,
                            "retries": retry_count
                        })
                        break
            
            # Move to next batch
            current_batch_start = batch_end
            
            # Small delay between batches to be respectful to the API
            if current_batch_start < total_chunks:
                time.sleep(1)
    
    except Exception as e:
        print(f"DEBUG: Error in batch processing: {str(e)}")
    
    result = {
        "total_chunks": total_chunks,
        "successful_chunks": successful_chunks,
        "failed_batches": failed_batches,
        "success_rate": (successful_chunks / total_chunks) * 100 if total_chunks > 0 else 0
    }
    
    print(f"DEBUG: Batch processing completed. Success: {successful_chunks}/{total_chunks} chunks "
          f"({result['success_rate']:.1f}%)")
    
    return result


def retry_failed_document_upload(progress_file_path: str, vector_store=None, batch_size=2, max_retries=5):
    """
    Retry processing failed chunks from a previous upload attempt.
    
    Args:
        progress_file_path: Path to the progress file from previous attempt
        vector_store: PineconeVectorStore instance (will create if None)
        batch_size: Number of chunks to process at once
        max_retries: Maximum number of retries for each batch
        
    Returns:
        dict: Results of retry attempt
    """
    try:
        # Load progress file
        if not os.path.exists(progress_file_path):
            return {"error": "Progress file not found"}
        
        with open(progress_file_path, 'r') as f:
            progress_data = json.load(f)
        
        if vector_store is None:
            vector_store = get_pinecone_vector_store()
        
        print(f"DEBUG: Resuming upload from chunk {progress_data.get('processed_chunks', 0)}")
        
        # This function would need the original document data to retry
        # For now, return information about what would need to be retried
        return {
            "message": "Resume functionality requires original document data",
            "progress_info": progress_data
        }
        
    except Exception as e:
        return {"error": f"Failed to retry upload: {str(e)}"}


def get_upload_progress_files():
    """
    Get list of progress files from failed uploads that can be resumed.
    
    Returns:
        list: List of progress files with their details
    """
    try:
        upload_dir = "uploads"
        if not os.path.exists(upload_dir):
            return []
        
        progress_files = []
        for filename in os.listdir(upload_dir):
            if filename.startswith("batch_progress_") and filename.endswith(".json"):
                file_path = os.path.join(upload_dir, filename)
                try:
                    with open(file_path, 'r') as f:
                        progress_data = json.load(f)
                    
                    progress_files.append({
                        "filename": filename,
                        "path": file_path,
                        "total_chunks": progress_data.get("total_chunks", 0),
                        "processed_chunks": progress_data.get("processed_chunks", 0),
                        "successful_chunks": progress_data.get("successful_chunks", 0),
                        "timestamp": progress_data.get("timestamp", "Unknown"),
                        "completion_rate": (progress_data.get("successful_chunks", 0) / 
                                          progress_data.get("total_chunks", 1)) * 100
                    })
                except:
                    continue
        
        return progress_files
        
    except Exception as e:
        print(f"DEBUG: Error getting progress files: {str(e)}")
        return []


def update_batch_upload_settings(batch_size: int = 3, max_retries: int = 5) -> dict:
    """
    Update batch upload settings in MongoDB config collection.
    
    Args:
        batch_size: Number of chunks to process in each batch
        max_retries: Maximum number of retries for failed batches
        
    Returns:
        dict: Status of the update operation
    """
    try:
        db = get_mongo_client()
        config_collection = db["config"]
        
        # Update batch size
        config_collection.update_one(
            {"key": "upload_batch_size"},
            {"$set": {"key": "upload_batch_size", "value": batch_size, "updated_at": datetime.now().isoformat()}},
            upsert=True
        )
        
        # Update max retries
        config_collection.update_one(
            {"key": "upload_max_retries"},
            {"$set": {"key": "upload_max_retries", "value": max_retries, "updated_at": datetime.now().isoformat()}},
            upsert=True
        )
        
        return {
            "success": True,
            "message": f"Batch upload settings updated: batch_size={batch_size}, max_retries={max_retries}"
        }
        
    except Exception as e:
        return {
            "success": False,
            "message": f"Failed to update batch settings: {str(e)}"
        }


def get_batch_upload_settings() -> dict:
    """
    Get current batch upload settings from MongoDB config collection.
    
    Returns:
        dict: Current batch upload settings
    """
    try:
        db = get_mongo_client()
        config_collection = db["config"]
        
        # Get batch size
        batch_size_config = config_collection.find_one({"key": "upload_batch_size"})
        batch_size = batch_size_config.get("value", 3) if batch_size_config else 3
        
        # Get max retries
        max_retries_config = config_collection.find_one({"key": "upload_max_retries"})
        max_retries = max_retries_config.get("value", 5) if max_retries_config else 5
        
        return {
            "batch_size": batch_size,
            "max_retries": max_retries,
            "success": True
        }
        
    except Exception as e:
        return {
            "batch_size": 3,
            "max_retries": 5,
            "success": False,
            "error": str(e)
        }


def process_uploaded_file(file: UploadFile, upload_type: str = "document") -> JSONResponse:
    """
    Process uploaded file: temporarily save, load content, add to vector store or database, then delete
    
    Args:
        file: UploadFile object from FastAPI
        upload_type: Type of upload - "document", "images_csv", "videos_csv"
        
    Returns:
        JSONResponse with success or error message
    """
    filename = f"{uuid.uuid4().hex}_{file.filename}"
    file_path = os.path.join("uploads", filename)

    try:
        # Save file temporarily
        with open(file_path, "wb") as f:
            f.write(file.file.read())

        # Handle different upload types
        if upload_type == "document":
            # Handle document files for knowledge base
            if file.filename.endswith(".pdf"):
                # For documents, don't delete file here - async processing will handle it
                return process_document_file(file_path, file.filename, "pdf")
            elif file.filename.endswith(".txt"):
                return process_document_file(file_path, file.filename, "txt")
            elif file.filename.endswith(".docx"):
                return process_document_file(file_path, file.filename, "docx")
            else:
                # Delete file for unsupported types
                if os.path.exists(file_path):
                    os.remove(file_path)
                return JSONResponse(content={"message": "❌ For document upload, please upload a PDF, TXT, or DOCX file."})
        
        elif upload_type == "images_csv":
            # Handle images CSV upload
            if file.filename.endswith(".csv"):
                result = upload_images_csv(file_path, file.filename)
                # Delete file after processing for CSV uploads
                if os.path.exists(file_path):
                    os.remove(file_path)
                return result
            else:
                # Delete file for unsupported types
                if os.path.exists(file_path):
                    os.remove(file_path)
                return JSONResponse(content={"message": "❌ For images upload, please upload a CSV file."})
        
        elif upload_type == "videos_csv":
            # Handle videos CSV upload
            if file.filename.endswith(".csv"):
                result = upload_videos_csv(file_path, file.filename)
                # Delete file after processing for CSV uploads
                if os.path.exists(file_path):
                    os.remove(file_path)
                return result
            else:
                # Delete file for unsupported types
                if os.path.exists(file_path):
                    os.remove(file_path)
                return JSONResponse(content={"message": "❌ For videos upload, please upload a CSV file."})
        
        else:
            # Delete file for invalid upload types
            if os.path.exists(file_path):
                os.remove(file_path)
            return JSONResponse(content={"message": "❌ Invalid upload type specified."})

    except Exception as e:
        # Delete file on error
        if os.path.exists(file_path):
            os.remove(file_path)
        return JSONResponse(content={"message": f"❌ Error processing file: {str(e)}"})


# Convenience functions for specific upload types
def process_document_upload(file: UploadFile) -> JSONResponse:
    """Process document file for knowledge base."""
    return process_uploaded_file(file, "document")


def process_images_csv_upload(file: UploadFile) -> JSONResponse:
    """Process images CSV file for images collection."""
    return process_uploaded_file(file, "images_csv")


def process_videos_csv_upload(file: UploadFile) -> JSONResponse:
    """Process videos CSV file for videos collection."""
    return process_uploaded_file(file, "videos_csv")


def process_document_file(file_path: str, filename: str, file_type: str) -> JSONResponse:
    """Process document files (PDF, TXT, DOCX) for vector store with batch processing."""
    try:
        if file_type == "pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
        elif file_type == "txt":
            print("DEBUG: Got text file, using simple Python file reading")
            # Read text file directly with Python instead of LangChain loader
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create a document-like object to maintain consistency
            docs = [type('Document', (), {
                'page_content': text_content,
                'metadata': {"source": filename}
            })()]
            print("DEBUG: Successfully read text file contents")
        elif file_type == "docx":
            # Use custom function to preserve layout and hyperlinks
            docx_content = extract_docx_with_layout_preserved(file_path)
            # Create a document-like object to maintain consistency
            docs = [type('Document', (), {
                'page_content': docx_content,
                'metadata': {"source": filename}
            })()]
        else:
            return JSONResponse(content={"message": "❌ Unsupported document type."})

        # Store in Pinecone with batch processing
        vector_store = get_pinecone_vector_store()
        
        # Process with batching (configurable batch size, default 3 chunks at a time)
        batch_size = 1  # Default value
        max_retries = 5  # Default value
        
        # Check if there's a custom batch size and max retries configuration
        try:
            db = get_mongo_client()
            config_collection = db["config"]
            batch_config = config_collection.find_one({"key": "upload_batch_size"})
            retries_config = config_collection.find_one({"key": "upload_max_retries"})
            
            if batch_config and batch_config.get("value"):
                batch_size = int(batch_config["value"])
                print(f"DEBUG: Using custom batch size from config: {batch_size}")
                
            if retries_config and retries_config.get("value"):
                max_retries = int(retries_config["value"])
                print(f"DEBUG: Using custom max retries from config: {max_retries}")
        except:
            print(f"DEBUG: Using default settings - batch size: {batch_size}, max retries: {max_retries}")
        
        print("DEBUG: Starting batch processing for document upload")
        batch_result = add_documents_to_vector_store_with_batching(
            vector_store,
            [doc.page_content for doc in docs],
            metadatas=[{"source": filename, "upload_time": datetime.now().isoformat()} for _ in docs],
            batch_size=batch_size,
            max_retries=max_retries
        )

        # Update last upload time only if at least some chunks were successful
        if batch_result["successful_chunks"] > 0:
            update_last_upload_time()
            
            # Add file record to MongoDB for tracking
            try:
                add_file_to_database(filename, datetime.now().isoformat(), "document")
            except Exception as db_error:
                print(f"DEBUG: Failed to add file to database: {str(db_error)}")
                # Don't fail the upload if database record fails

        # Prepare response message based on results
        if batch_result["failed_batches"]:
            if batch_result["successful_chunks"] > 0:
                # Partial success
                message = (f"⚠️ Partially processed {filename}. "
                          f"Successfully added {batch_result['successful_chunks']} chunks "
                          f"out of {batch_result['total_chunks']} to Ask Nour's knowledge base. "
                          f"{len(batch_result['failed_batches'])} batches failed due to API limits. "
                          f"You may try uploading again to process the remaining chunks.")
            else:
                # Complete failure
                message = (f"❌ Failed to process {filename}. "
                          f"All {batch_result['total_chunks']} chunks failed due to API quota limits. "
                          f"Please try again later when your API quota resets.")
        else:
            # Complete success
            message = (f"✅ Successfully processed {filename} and added "
                      f"{batch_result['total_chunks']} chunks to Ask Nour's knowledge base!")

        return JSONResponse(content={
            "message": message,
            "details": {
                "total_chunks": batch_result["total_chunks"],
                "successful_chunks": batch_result["successful_chunks"],
                "failed_batches": len(batch_result["failed_batches"]),
                "success_rate": f"{batch_result['success_rate']:.1f}%"
            }
        })

    except Exception as e:
        return JSONResponse(content={"message": f"❌ Error processing document: {str(e)}"})
    
    finally:
        # Always clean up the temporary file after processing
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"DEBUG: Cleaned up temporary file: {file_path}")
        except Exception as cleanup_error:
            print(f"DEBUG: Error cleaning up file {file_path}: {cleanup_error}")


def upload_images_csv(file_path: str, filename: str) -> JSONResponse:
    """Upload images from CSV to MongoDB images collection with hardcoded keys."""
    try:
        db = get_mongo_client()
        collection = db[IMAGES_COLLECTION]
        
        with open(file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            
            # Skip header row (first row contains column headers and is ignored)
            next(reader, None)
            
            images = []
            for row_num, row in enumerate(reader, 1):
                if len(row) >= 1:  # At least image URL required
                    image_doc = {
                        "image_url": row[0].strip(),
                        "image_description": row[1].strip() if len(row) > 1 else "",
                        "uploaded_at": datetime.now().isoformat(),
                        "uploaded_from": filename,
                        "row_number": row_num
                    }
                    images.append(image_doc)
            
            if images:
                collection.insert_many(images)
                # Update last upload time
                update_last_upload_time()
                
                # Add file record to tracking database
                try:
                    add_file_to_database(filename, datetime.now().isoformat(), "images_csv")
                except Exception as db_error:
                    print(f"DEBUG: Failed to add images CSV file to database: {str(db_error)}")
                
                return JSONResponse(content={
                    "message": f"✅ Successfully uploaded {len(images)} images from {filename} to the images collection!",
                    "count": len(images)
                })
            else:
                return JSONResponse(content={"message": "❌ No valid image data found in CSV file."})
                
    except Exception as e:
        return JSONResponse(content={"message": f"❌ Error uploading images CSV: {str(e)}"})


def upload_videos_csv(file_path: str, filename: str) -> JSONResponse:
    """Upload videos from CSV to MongoDB videos collection with hardcoded keys."""
    try:
        db = get_mongo_client()
        collection = db[VIDEOS_COLLECTION]
        
        with open(file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            
            # Skip header row (first row contains column headers and is ignored)
            next(reader, None)
            
            videos = []
            for row_num, row in enumerate(reader, 1):
                if len(row) >= 1:  # At least video URL required
                    video_doc = {
                        "video_url": row[0].strip(),
                        "video_description": row[1].strip() if len(row) > 1 else "",
                        "uploaded_at": datetime.now().isoformat(),
                        "uploaded_from": filename,
                        "row_number": row_num
                    }
                    videos.append(video_doc)
            
            if videos:
                collection.insert_many(videos)
                # Update last upload time
                update_last_upload_time()
                
                # Add file record to tracking database
                try:
                    add_file_to_database(filename, datetime.now().isoformat(), "videos_csv")
                except Exception as db_error:
                    print(f"DEBUG: Failed to add videos CSV file to database: {str(db_error)}")
                
                return JSONResponse(content={
                    "message": f"✅ Successfully uploaded {len(videos)} videos from {filename} to the videos collection!",
                    "count": len(videos)
                })
            else:
                return JSONResponse(content={"message": "❌ No valid video data found in CSV file."})
                
    except Exception as e:
        return JSONResponse(content={"message": f"❌ Error uploading videos CSV: {str(e)}"})


def update_last_upload_time():
    """Update the last upload timestamp in MongoDB extras collection"""
    try:
        db = get_mongo_client()
        collection = db[EXTRAS_COLLECTION]
        
        # Upsert the last upload time document
        collection.update_one(
            {"type": "last_upload"},
            {
                "$set": {
                    "type": "last_upload",
                    "timestamp": datetime.now().isoformat(),
                    "updated_at": datetime.now().isoformat()
                }
            },
            upsert=True
        )
    except Exception as e:
        print(f"DEBUG: Error updating last upload time: {str(e)}")


def get_last_upload_time():
    """Get the last upload timestamp from MongoDB extras collection"""
    try:
        db = get_mongo_client()
        collection = db[EXTRAS_COLLECTION]
        
        # Find the last upload document
        last_upload_doc = collection.find_one({"type": "last_upload"})
        
        if last_upload_doc and "timestamp" in last_upload_doc:
            timestamp = last_upload_doc["timestamp"]
            dt = datetime.fromisoformat(timestamp)
            return dt.strftime("%B %d, %Y at %I:%M %p")
    except Exception as e:
        print(f"DEBUG: Error getting last upload time: {str(e)}")
    
    return "Never"


def get_knowledge_base_stats():
    """Get statistics about the knowledge base from Pinecone and MongoDB collections"""
    try:
        # Get Pinecone stats
        pinecone_stats = get_pinecone_stats()
        
        # Get MongoDB stats
        db = get_mongo_client()
        images_count = db[IMAGES_COLLECTION].count_documents({})
        videos_count = db[VIDEOS_COLLECTION].count_documents({})
        
        return {
            "total_vectors": pinecone_stats.get('total_vector_count', 0),
            "total_images": images_count,
            "total_videos": videos_count,
            "last_upload": get_last_upload_time()
        }
    except Exception as e:
        print(f"DEBUG: Error getting knowledge base stats: {str(e)}")
        return {
            "total_vectors": 0,
            "total_images": 0,
            "total_videos": 0,
            "last_upload": "Never"
        }


def get_uploaded_files_list():
    """Get list of uploaded files in the knowledge base"""
    try:
        return list_uploaded_files()
    except Exception as e:
        print(f"DEBUG: Error getting uploaded files: {str(e)}")
        return []


def delete_uploaded_file(filename):
    """Delete an uploaded file from the knowledge base"""
    try:
        result = delete_file_from_pinecone(filename)
        
        if result["success"]:
            # Update last upload time after deletion
            update_last_upload_time()
        
        return result
    except Exception as e:
        print(f"DEBUG: Error deleting uploaded file: {str(e)}")
        return {"success": False, "message": f"Error deleting file: {str(e)}"}


# This function is to be used insdie extract_docx_with_layout_preserved ...
def extract_hyperlinks_from_paragraph(paragraph):
    """
    Extract hyperlinks from a paragraph by parsing the XML structure.
    """
    hyperlinks = {}
    
    # Get all hyperlink elements in the paragraph
    for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
        r_id = hyperlink.get(qn('r:id'))
        if r_id:
            # Get the URL from the document relationships
            try:
                url = paragraph.part.rels[r_id].target_ref
                # Get the text content of the hyperlink
                link_text = ''.join([node.text for node in hyperlink.xpath('.//w:t')])
                hyperlinks[link_text] = url
            except KeyError:
                continue
    
    return hyperlinks

# this function is only to be used with docx files to preserve layout and hyperlinks
def extract_docx_with_layout_preserved(file_path):
    """
    Extract text content from DOCX while preserving layout and hyperlinks.
    Hyperlinks appear exactly where they are in the document.
    Tables appear in their original positions.
    """
    doc = Document(file_path)
    content_parts = []
    
    # Keep track of table index
    table_index = 0
    
    # Process document elements in their original order
    for element in doc.element.body:
        # Check if element is a paragraph
        if element.tag.endswith('p'):
            # Find the corresponding paragraph object
            for paragraph in doc.paragraphs:
                if paragraph._element == element:
                    if not paragraph.text.strip():
                        # Preserve empty lines for layout
                        content_parts.append("")
                        continue
                    
                    # Get hyperlinks for this paragraph
                    hyperlinks = extract_hyperlinks_from_paragraph(paragraph)
                    
                    # Build paragraph text with hyperlinks
                    para_text = paragraph.text
                    
                    # Replace hyperlink text with text + URL format
                    for link_text, url in hyperlinks.items():
                        if link_text in para_text:
                            para_text = para_text.replace(link_text, f"{link_text} [{url}]")
                    
                    content_parts.append(para_text)
                    break
        
        # Check if element is a table
        elif element.tag.endswith('tbl'):
            content_parts.append(f"\n [TABLE {table_index + 1}]")
            table = doc.tables[table_index]
            
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            # Get hyperlinks for this cell paragraph
                            hyperlinks = extract_hyperlinks_from_paragraph(paragraph)
                            
                            para_text = paragraph.text
                            # Replace hyperlink text with text + URL format
                            for link_text, url in hyperlinks.items():
                                if link_text in para_text:
                                    para_text = para_text.replace(link_text, f"{link_text} [{url}]")
                            
                            cell_text += para_text + " "
                    
                    row_content.append(cell_text.strip())
                
                # Format table row with proper spacing
                content_parts.append(" | ".join(row_content))
            
            content_parts.append("[END TABLE]\n ")
            table_index += 1
    
    return "\n ".join(content_parts)
